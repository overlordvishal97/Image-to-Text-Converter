import os
import docx
from PIL import Image
import pytesseract
import warnings

# Ignore the Warning
warnings.filterwarnings("ignore", category=Warning)

# Set the path to the Tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


# Set the input and output directories
input_dir = r'C:\Users\visha\Downloads\Assignment Files'
output_dir = r'C:\Users\visha\Downloads\wfh2'

# Make sure the output directory exists
os.makedirs(output_dir, exist_ok=True)

# Iterate through all image files in the input directory
for filename in os.listdir(input_dir):
    if filename.endswith(('.png', '.jpg', '.jpeg')):
        img_path = os.path.join(input_dir, filename)
        img = Image.open(img_path)
        text = pytesseract.image_to_string(img)

        doc_path = os.path.join(output_dir, filename[:-4] + '.docx')
        doc = docx.Document()
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = docx.shared.Pt(12)
        doc.add_paragraph(text)
        doc.save(doc_path)

        # Extract text from the DOCX file
        doc = docx.Document(doc_path)
        docx_text = ''
        for para in doc.paragraphs:
            docx_text += para.text

        # Compare the extracted text with the text in the DOCX file
        if text == docx_text:
            print(f'Text in {filename} matches text in {os.path.basename(doc_path)}')
        else:
            print(f'Text in {filename} does not match text in {os.path.basename(doc_path)}')

        os.remove(img_path)

print('All image files have been processed.')